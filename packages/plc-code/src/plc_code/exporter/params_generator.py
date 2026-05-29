"""Parameter table document generator.

Generates a Word document (.docx) containing organized parameter tables
from parsed TIA Portal InstanceDB XML exports. Supports template-based
cover pages and smart unit filtering.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips

from plc_code.exporter.param_config import FilterRule, ParamsExportConfig
from plc_code.exporter.param_parser import (
    ParameterEntry,
    get_arm_names,
    get_arm_types,
    group_entries,
    parse_parameter_xml,
)


@dataclass
class GenerationResult:
    """Result of document generation.

    Attributes
    ----------
    success : bool
        Whether generation succeeded.
    output_path : Path | None
        Path to the generated document.
    error : str
        Error message if generation failed.
    entry_count : int
        Total number of parameter entries in the document.
    """

    success: bool
    output_path: Path | None = None
    error: str = ""
    entry_count: int = 0


# Table column widths in twips (matching the template)
COL_WIDTHS = [Twips(6233), Twips(1558), Twips(1419), Twips(5918)]
HEADER_LABELS = ["Qualified Path", "Data Type", "Value", "Comment"]

# Unit parameter sub-sections (relative path prefix after arms[N].)
ARM_SUBSECTIONS = [
    ("Motion — Angular", "motion.angular."),
    ("Motion — Retraction", "motion.retraction."),
    ("Geometry", "geometry."),
    ("Limits", "limit."),
    ("CPMS", "cpms["),
]


def generate_params_document(config: ParamsExportConfig) -> GenerationResult:
    """Generate the parameter table document.

    Parameters
    ----------
    config : ParamsExportConfig
        Export configuration.

    Returns
    -------
    GenerationResult
        Result of the generation.
    """
    # Parse all XML sources
    all_entries: list[ParameterEntry] = []
    source_groups: dict[str, list[ParameterEntry]] = {}

    for source in config.sources:
        source_path = config.resolve_source_path(source)
        types_dir = (config.root_path / source.types_dir) if source.types_dir else None
        try:
            entries = parse_parameter_xml(source_path, source.prefix, types_dir=types_dir)
        except (FileNotFoundError, ValueError) as e:
            return GenerationResult(success=False, error=str(e))

        all_entries.extend(entries)
        source_groups[source.prefix] = entries

    if not all_entries:
        return GenerationResult(success=False, error="No parameter entries found in XML sources")

    # Apply descriptions from config
    for entry in all_entries:
        if not entry.comment:
            # Get relative path (strip db prefix)
            for source in config.sources:
                if entry.qualified_path.startswith(source.prefix + "."):
                    rel = entry.qualified_path[len(source.prefix) + 1 :]
                    desc = config.get_description(rel)
                    if desc:
                        entry.comment = desc
                    break

    # Apply conditional filters
    if config.filters:
        for source in config.sources:
            if source.prefix in source_groups:
                source_groups[source.prefix] = _apply_filters(
                    source_groups[source.prefix], config.filters, source.prefix
                )
        # Rebuild all_entries from filtered groups
        all_entries = []
        for source in config.sources:
            all_entries.extend(source_groups.get(source.prefix, []))

    # Determine arm types and names from parsed data
    arm_types: dict[int, str] = {}
    arm_names: dict[int, str] = {}
    for source in config.sources:
        if source.prefix in source_groups:
            arm_types.update(get_arm_types(source_groups[source.prefix], source.prefix))
            arm_names.update(get_arm_names(source_groups[source.prefix], source.prefix))

    # Open template or create new document
    if config.template_path and config.template_path.exists():
        doc = Document(str(config.template_path))
        _fill_cover_page(doc, config)
        _remove_existing_parameter_tables(doc)
    else:
        doc = Document()
        _add_landscape_section(doc)
        doc.add_heading("Introduction", level=2)
        doc.add_paragraph("This document exposes all the parameters of the program.")

    # Build organized content
    total_entries = 0

    for source in config.sources:
        prefix = source.prefix
        entries = source_groups.get(prefix, [])
        if not entries:
            continue

        groups = group_entries(entries, prefix)

        # Add source heading
        doc.add_heading(f"{prefix}", level=2)

        # Global parameters
        if "global" in groups:
            doc.add_heading("Global Parameters", level=3)
            total_entries += _add_parameter_table(doc, groups["global"])

        # Unit parameters — iterate over all unit groups found in the data
        arm_indices = sorted(int(k.split(".")[1]) for k in groups if k.startswith("arms."))
        for arm_idx in arm_indices:
            arm_key = f"arms.{arm_idx}"
            arm_entries = groups[arm_key]
            arm_type = arm_types.get(arm_idx, "")
            type_clean = arm_type.strip("'\"")

            # Resolve label: config override → XML name field → fallback
            label = _resolve_arm_label(arm_idx, config.arms.labels, arm_names)

            # Skip inactive arms
            if config.arms.skip_inactive and type_clean == "none":
                # Still show the type entry as a single row
                type_entries = [e for e in arm_entries if e.qualified_path.endswith(".type")]
                if type_entries:
                    doc.add_heading(f"{label} (type: {type_clean})", level=3)
                    p = doc.add_paragraph("Unit is inactive — parameters skipped.")
                    p.style = doc.styles["Normal"]
                    total_entries += 1
                continue

            doc.add_heading(f"{label} (type: {type_clean})", level=3)

            # Group unit entries into sub-sections
            remaining = list(arm_entries)
            arm_prefix = f"{prefix}.arms[{arm_idx}]."

            for subsection_name, subsection_path in ARM_SUBSECTIONS:
                sub_entries = []
                for entry in remaining:
                    rel = entry.qualified_path[len(arm_prefix) :]
                    if rel.startswith(subsection_path):
                        sub_entries.append(entry)

                if sub_entries:
                    doc.add_heading(subsection_name, level=4)
                    total_entries += _add_parameter_table(doc, sub_entries)
                    for entry in sub_entries:
                        remaining.remove(entry)

            # Remaining unit entries (type, etc.) that don't fit a subsection
            if remaining:
                doc.add_heading("General", level=4)
                total_entries += _add_parameter_table(doc, remaining)

        # Safety unit parameters
        safety_arm_indices = sorted(int(k.split(".")[1]) for k in groups if k.startswith("safety_arms."))
        for arm_idx in safety_arm_indices:
            arm_key = f"safety_arms.{arm_idx}"
            label = _resolve_arm_label(arm_idx, config.arms.labels, arm_names)
            doc.add_heading(f"Safety — {label}", level=3)
            total_entries += _add_parameter_table(doc, groups[arm_key])

        # Redundancy parameters
        if "redundancy" in groups:
            doc.add_heading("Redundancy Parameters", level=3)
            total_entries += _add_parameter_table(doc, groups["redundancy"])

        # Remote parameters
        if "remote" in groups:
            doc.add_heading("Remote Parameters", level=3)
            total_entries += _add_parameter_table(doc, groups["remote"])

        # Drive parameters
        if "drive" in groups:
            drive_entries = groups["drive"]
            doc.add_heading("Drive Parameters", level=3)

            # Sub-group drive entries
            scaling_entries = [
                e
                for e in drive_entries
                if "Scaling" in e.qualified_path or "scaling" in e.qualified_path.lower()
            ]
            timing_entries = [e for e in drive_entries if e.datatype == "Time" and e not in scaling_entries]
            threshold_entries = [
                e for e in drive_entries if e.datatype == "Real" and e not in scaling_entries
            ]
            general_entries = [
                e
                for e in drive_entries
                if e not in scaling_entries and e not in timing_entries and e not in threshold_entries
            ]

            if general_entries:
                doc.add_heading("General", level=4)
                total_entries += _add_parameter_table(doc, general_entries)
            if scaling_entries:
                doc.add_heading("Transmitter Scaling", level=4)
                total_entries += _add_parameter_table(doc, scaling_entries)
            if timing_entries:
                doc.add_heading("Timing Delays", level=4)
                total_entries += _add_parameter_table(doc, timing_entries)
            if threshold_entries:
                doc.add_heading("Pressure & Temperature Thresholds", level=4)
                total_entries += _add_parameter_table(doc, threshold_entries)

        # ControlUnit parameters
        if "controller" in groups:
            doc.add_heading("ControlUnit Parameters", level=3)
            total_entries += _add_parameter_table(doc, groups["controller"])

        # Modbus parameters
        if "modbus" in groups:
            doc.add_heading("Modbus Parameters", level=3)
            total_entries += _add_parameter_table(doc, groups["modbus"])

    # Save document
    output_path = config.output_path
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return GenerationResult(
        success=True,
        output_path=output_path,
        entry_count=total_entries,
    )


def _resolve_arm_label(
    arm_idx: int,
    config_labels: dict[int, str],
    arm_names: dict[int, str],
) -> str:
    """Resolve the display label for a unit.

    Priority: config override → XML ``name`` field → fallback ``Unit N``.

    Parameters
    ----------
    arm_idx : int
        Unit array index.
    config_labels : dict[int, str]
        Labels from configuration (may be empty).
    arm_names : dict[int, str]
        Names extracted from the ``arms[N].name`` parameter in the XML.

    Returns
    -------
    str
        Human-readable unit label.
    """
    if arm_idx in config_labels:
        return config_labels[arm_idx]
    raw_name = arm_names.get(arm_idx, "")
    clean = raw_name.strip("'\"")
    if clean:
        return clean
    return f"Unit {arm_idx}"


def _apply_filters(
    entries: list[ParameterEntry],
    filters: list[FilterRule],
    db_name: str,
) -> list[ParameterEntry]:
    """Apply conditional filter rules to remove entries.

    Each filter rule defines a scope pattern (e.g. "arms[*].cpms[*]"),
    a condition (field + value), and an action (skip_section or skip specific fields).

    Parameters
    ----------
    entries : list[ParameterEntry]
        All parameter entries for a source.
    filters : list[FilterRule]
        Filter rules from config.
    db_name : str
        DB name prefix (e.g. "ProcessParameter").

    Returns
    -------
    list[ParameterEntry]
        Filtered entries with matching entries removed.
    """
    to_remove: set[int] = set()  # indices of entries to remove

    for rule in filters:
        if not rule.scope or not rule.field:
            continue

        # Convert scope pattern to regex: "arms[*].cpms[*]" → "arms\[\d+\]\.cpms\[\d+\]"
        scope_regex = re.escape(rule.scope).replace(r"\[\*\]", r"\[(\d+)\]")
        # Match entries AT the scope level (e.g. the scope itself as prefix)
        scope_prefix_pattern = re.compile(rf"^{re.escape(db_name)}\.{scope_regex}(?:\.|$)")

        # Find all unique scope instances
        scope_instances: dict[str, list[int]] = {}
        for idx, entry in enumerate(entries):
            m = scope_prefix_pattern.match(entry.qualified_path)
            if m:
                # Extract the concrete scope prefix (e.g. "ProcessParameter.arms[1].cpms[0]")
                scope_concrete = entry.qualified_path[: m.end()].rstrip(".")
                scope_instances.setdefault(scope_concrete, []).append(idx)

        # Evaluate condition for each scope instance
        for scope_prefix, entry_indices in scope_instances.items():
            # Find the condition field within this scope
            condition_path = f"{scope_prefix}.{rule.field}"
            condition_value = None
            for idx in entry_indices:
                if entries[idx].qualified_path == condition_path:
                    condition_value = entries[idx].value
                    break

            if condition_value is None or condition_value != rule.equals:
                continue

            # Condition matched — apply action
            if rule.action == "skip_section":
                to_remove.update(entry_indices)
            elif rule.skip:
                for idx in entry_indices:
                    rel = entries[idx].qualified_path[len(scope_prefix) + 1 :]
                    for skip_prefix in rule.skip:
                        if rel.startswith(skip_prefix):
                            to_remove.add(idx)
                            break

    if not to_remove:
        return entries

    return [e for i, e in enumerate(entries) if i not in to_remove]


def _add_parameter_table(doc: Any, entries: list[ParameterEntry]) -> int:
    """Add a parameter table to the document.

    Parameters
    ----------
    doc : Document
        The docx document.
    entries : list[ParameterEntry]
        Parameter entries to add as rows.

    Returns
    -------
    int
        Number of entries added.
    """
    if not entries:
        return 0

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style the table with borders
    _set_table_borders(table)

    # Header row
    header_cells = table.rows[0].cells
    for i, label in enumerate(HEADER_LABELS):
        p = header_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)

    # Set column widths
    for i, width in enumerate(COL_WIDTHS):
        header_cells[i].width = width

    # Data rows
    for entry in entries:
        row = table.add_row()
        cells = row.cells

        values = [
            entry.qualified_path,
            entry.datatype,
            entry.value,
            entry.comment,
        ]
        for i, val in enumerate(values):
            p = cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            cells[i].width = COL_WIDTHS[i]

    # Add spacing after table
    doc.add_paragraph("")

    return len(entries)


def _set_table_borders(table: object) -> None:
    """Apply grid borders to a table.

    Parameters
    ----------
    table : Table
        The docx table object.
    """
    tbl = table._tbl  # type: ignore[attr-defined]
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)

    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def _fill_cover_page(doc: Any, config: ParamsExportConfig) -> None:
    """Fill in the cover page template fields.

    Parameters
    ----------
    doc : Document
        The docx document (opened from template).
    config : ParamsExportConfig
        Export configuration with project metadata.
    """
    # Replace MERGEFIELD placeholders in all paragraphs
    field_values = {
        "Tag_number": config.project.tag_number,
        "Drawing_Title": config.document.title,
        "Rev": str(config.revision.rev),
    }
    _replace_merge_fields(doc, field_values)

    # Update cover table (Table 0) if it exists
    if len(doc.tables) < 1:
        return

    cover = doc.tables[0]

    # Row 1: tag number and customer
    if len(cover.rows) > 1:
        row = cover.rows[1]
        if config.project.tag_number:
            _set_cell_text(row.cells[0], config.project.tag_number)
        if config.project.customer:
            # Customer spans cells 2-6
            _set_cell_text(row.cells[2], config.project.customer)

    # Row 2: equipment description
    if len(cover.rows) > 2 and config.project.equipment:
        _set_cell_text(cover.rows[2].cells[0], config.project.equipment)

    # Row 3: document title
    if len(cover.rows) > 3 and config.document.title:
        _set_cell_text(cover.rows[3].cells[0], f"\n\n\n{config.document.title}\n\n")

    # Revision rows (rows 8, 9, 10 in the template)
    # The template has revision rows above the header row (row 11)
    # We fill the first revision row with current revision
    all_revisions = list(config.history) if config.history else []
    if not all_revisions or all_revisions[0].rev != config.revision.rev:
        all_revisions.insert(0, config.revision)

    # Find revision rows (between row 7 and the "Rev." header row)
    rev_header_idx = None
    for i, row in enumerate(cover.rows):
        if any(cell.text.strip() == "Rev." for cell in row.cells):
            rev_header_idx = i
            break

    if rev_header_idx is not None:
        # Revision rows are above the header, in descending order
        for rev_offset, rev_entry in enumerate(all_revisions):
            row_idx = rev_header_idx - 1 - rev_offset
            if row_idx < 0 or row_idx >= len(cover.rows):
                break
            row = cover.rows[row_idx]
            cells = row.cells
            if len(cells) >= 9:
                _set_cell_text(cells[0], str(rev_entry.rev))
                _set_cell_text(cells[1], rev_entry.date)
                _set_cell_text(cells[3], rev_entry.description)
                _set_cell_text(cells[5], rev_entry.prepared_by)
                _set_cell_text(cells[6], rev_entry.checked_by)
                _set_cell_text(cells[8], rev_entry.approved_by)

    # Document number row (row after Rev. header)
    if rev_header_idx is not None and rev_header_idx + 1 < len(cover.rows):
        doc_row = cover.rows[rev_header_idx + 1]
        if config.document.number:
            _set_cell_text(doc_row.cells[3], config.document.number)


def _remove_existing_parameter_tables(doc: Any) -> None:
    """Remove existing parameter tables and content after the introduction.

    Keeps the cover page (section 0) and introduction text intact.
    Removes everything from the "parameters table" heading onwards.

    Parameters
    ----------
    doc : Document
        The docx document.
    """
    # Find the "parameters table" heading or similar
    body = doc.element.body
    remove_from = None

    for i, element in enumerate(body):
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            # Check paragraph text
            text = element.text or ""
            # Also check runs
            for run in element.iter():
                run_tag = run.tag.split("}")[-1] if "}" in run.tag else run.tag
                if run_tag == "t" and run.text:
                    text += run.text
            if "parameters table" in text.lower() or "parameter table" in text.lower():
                remove_from = i
                break

    if remove_from is not None:
        # Remove everything from the heading onwards (but keep the heading itself)
        elements_to_remove = list(body)[remove_from + 1 :]
        for elem in elements_to_remove:
            # Don't remove section properties
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "sectPr":
                continue
            body.remove(elem)


def _replace_merge_fields(doc: Any, values: dict[str, str]) -> None:
    """Replace MERGEFIELD placeholders with values.

    Parameters
    ----------
    doc : Document
        The docx document.
    values : dict[str, str]
        Field name → replacement value mapping.
    """
    # MERGEFIELD instructions are in complex field sequences:
    # <w:fldChar w:fldCharType="begin"/> ... <w:instrText> MERGEFIELD Name </w:instrText>
    # ... <w:fldChar w:fldCharType="separate"/> ... <w:t>display text</w:t>
    # ... <w:fldChar w:fldCharType="end"/>
    #
    # We find instrText elements, identify the field name, then replace
    # the display text between separate and end markers.
    body = doc.element.body
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    for paragraph in body.iter(f"{{{ns}}}p"):
        runs = list(paragraph.iter(f"{{{ns}}}r"))
        field_name = None
        past_separate = False

        for run in runs:
            # Check for instrText
            instr = run.find(f"{{{ns}}}instrText")
            if instr is not None and instr.text:
                for name in values:
                    if f"MERGEFIELD {name}" in instr.text or f"MERGEFIELD  {name}" in instr.text:
                        field_name = name
                        break

            # Check for fldChar
            fld_char = run.find(f"{{{ns}}}fldChar")
            if fld_char is not None:
                fld_type = fld_char.get(f"{{{ns}}}fldCharType")
                if fld_type == "begin":
                    past_separate = False
                elif fld_type == "separate":
                    past_separate = True
                elif fld_type == "end":
                    field_name = None
                    past_separate = False

            # Replace display text
            if past_separate and field_name and field_name in values:
                t_elem = run.find(f"{{{ns}}}t")
                if t_elem is not None:
                    t_elem.text = values[field_name]


def _set_cell_text(cell: Any, text: str) -> None:
    """Set the text of a table cell, preserving formatting of the first run.

    Clears all paragraphs except the first, and all runs except the first.

    Parameters
    ----------
    cell : Cell
        The docx table cell.
    text : str
        New text content.
    """
    if cell.paragraphs:
        # Set text in first paragraph, first run
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = text

        # Remove all additional paragraphs from the cell XML
        tc = cell._tc
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        paragraphs = list(tc.findall(f"{{{ns}}}p"))
        for extra_p in paragraphs[1:]:
            tc.remove(extra_p)
    else:
        cell.text = text


def _add_landscape_section(doc: Any) -> None:
    """Add a landscape section to the document.

    Parameters
    ----------
    doc : Document
        The docx document.
    """
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
